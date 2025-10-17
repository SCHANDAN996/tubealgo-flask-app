from flask import render_template, request, redirect, url_for, flash, Blueprint, Response
from flask_login import login_required, current_user
from tubealgo.models import Competitor
from tubealgo.services.youtube_fetcher import (
    analyze_channel, get_full_video_details, get_all_channel_videos,
    get_most_used_tags, get_most_viewed_videos, get_latest_videos,
    get_upload_schedule_analysis
)
from tubealgo.routes.api_routes import get_full_competitor_package
from tubealgo.routes.utils import get_video_info_dict, sanitize_filename
from datetime import datetime, timezone
import json
from openpyxl import Workbook
from docx import Document
from io import BytesIO
from urllib.parse import quote # <-- यह नया इम्पोर्ट है

analysis_bp = Blueprint('analysis', __name__)

@analysis_bp.route('/deep-analysis/<string:channel_id>')
@login_required
def deep_analysis(channel_id):
    competitor = Competitor.query.filter_by(user_id=current_user.id, channel_id_youtube=channel_id).first()
    if not competitor:
        flash("Competitor not found in your list.", "error")
        return redirect(url_for('competitor.competitors'))
    
    data_package = get_full_competitor_package(competitor.id)

    if 'error' in data_package:
        flash(data_package['error'], 'error')
        return redirect(url_for('competitor.competitors'))

    channel_data = data_package.get('details', {})
    
    recent_videos_data = data_package.get('recent_videos_data', {}) or {}
    most_viewed_data = data_package.get('most_viewed_videos_data', {}) or {}
    
    recent_videos_list = recent_videos_data.get('videos', [])
    most_viewed_videos_list = most_viewed_data.get('videos', [])

    top_tags = data_package.get('top_tags', [])
    playlists = data_package.get('playlists', [])
    
    upload_schedule = get_upload_schedule_analysis(channel_id)
    channel_keywords = channel_data.get('keywords', [])

    upload_counts = {}
    
    if recent_videos_list:
        for video in recent_videos_list:
            if video.get('upload_date'):
                upload_month = datetime.fromisoformat(video['upload_date'].replace('Z', '+00:00')).strftime('%Y-%m')
                upload_counts[upload_month] = upload_counts.get(upload_month, 0) + 1
    
    sorted_months = sorted(upload_counts.keys(), reverse=True)[:6]
    sorted_months.reverse()
    upload_labels = [datetime.strptime(month, '%Y-%m').strftime('%b %Y') for month in sorted_months]
    upload_data = [upload_counts.get(month, 0) for month in sorted_months]

    all_videos_dict = {v['id']: v for v in (recent_videos_list + most_viewed_videos_list) if v and 'id' in v}
    all_videos_unique = list(all_videos_dict.values())
    
    processed_videos_for_export = []
    for v in all_videos_unique:
        processed_videos_for_export.append({
            'id': v.get('id'),
            'title': v.get('title'),
            'thumbnail': v.get('thumbnail'),
            'view_count': v.get('view_count', 0),
            'like_count': v.get('like_count', 0),
            'comment_count': v.get('comment_count', 0),
            'upload_date': v.get('upload_date'),
            'duration_seconds': v.get('duration_seconds', 0),
            'is_short': v.get('is_short', False)
        })

    return render_template('deep_analysis.html', 
                           channel_data=channel_data,
                           top_tags=top_tags,
                           playlists=playlists,
                           upload_labels=json.dumps(upload_labels),
                           upload_data=json.dumps(upload_data),
                           recent_videos_json=json.dumps(processed_videos_for_export),
                           most_viewed_videos_json=json.dumps([]),
                           channel_keywords=channel_keywords,
                           upload_schedule_json=json.dumps(upload_schedule),
                           active_page='competitors')


@analysis_bp.route('/video/<video_id>')
@login_required
def video_analysis(video_id):
    video_info = get_video_info_dict(video_id)
    if 'error' in video_info:
        flash(video_info['error'], 'error')
        return redirect(request.referrer or url_for('competitor.competitors'))
    return render_template('video_analysis.html', video=video_info)

@analysis_bp.route('/analysis/export/excel/<string:channel_id>')
@login_required
def export_channel_videos_to_excel(channel_id):
    selected_columns = request.args.getlist('columns')
    if not selected_columns:
        return "Please select at least one column to export.", 400

    all_videos = get_all_channel_videos(channel_id)
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Channel Video Data"

    headers = [col.replace('_', ' ').title() for col in selected_columns]
    ws.append(headers)

    for video in all_videos:
        row = []
        for col in selected_columns:
            if col == 'video_url':
                row.append(f"https://www.youtube.com/watch?v={video.get('id', '')}")
            elif col == 'upload_date':
                date_str = video.get(col, '')
                if date_str:
                    row.append(datetime.fromisoformat(date_str.replace('Z', '')).strftime('%Y-%m-%d %H:%M'))
                else:
                    row.append('')
            else:
                row.append(video.get(col, 'N/A'))
        ws.append(row)
    
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    channel_info = analyze_channel(channel_id)
    filename = f"{sanitize_filename(channel_info.get('Title', 'channel'))}_videos_export.xlsx"
    
    # --- यहाँ बदलाव किया गया है ---
    ascii_filename = filename.encode('ascii', 'ignore').decode('ascii', 'ignore')
    headers = {
        'Content-Disposition': 'attachment; filename*=UTF-8\'\'{}; filename="{}"'.format(quote(filename), ascii_filename)
    }
    
    return Response(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers=headers
    )

@analysis_bp.route('/analysis/export_video/excel/<string:video_id>')
@login_required
def export_single_video_to_excel(video_id):
    video_info = get_video_info_dict(video_id)
    if 'error' in video_info:
        flash(f"Could not export data: {video_info['error']}", "error")
        return redirect(url_for('analysis.video_analysis', video_id=video_id))

    wb = Workbook()
    ws = wb.active
    ws.title = "Video Analysis"
    
    ws.append(['Metric', 'Value'])
    for key, value in video_info.items():
        if key in ['tags', 'hashtags']:
            ws.append([key.title(), ', '.join(value)])
        elif not isinstance(value, dict):
            ws.append([key.replace('_', ' ').title(), value])

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    filename = f"{sanitize_filename(video_info.get('title', 'video'))}_analysis.xlsx"
    
    # --- यहाँ बदलाव किया गया है ---
    ascii_filename = filename.encode('ascii', 'ignore').decode('ascii', 'ignore')
    headers = {
        'Content-Disposition': 'attachment; filename*=UTF-8\'\'{}; filename="{}"'.format(quote(filename), ascii_filename)
    }
    
    return Response(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers=headers
    )

@analysis_bp.route('/analysis/export_video/word/<string:video_id>')
@login_required
def export_single_video_to_word(video_id):
    video_info = get_video_info_dict(video_id)
    if 'error' in video_info:
        flash(f"Could not export data: {video_info['error']}", "error")
        return redirect(url_for('analysis.video_analysis', video_id=video_id))

    document = Document()
    document.add_heading(video_info['title'], level=1)
    document.add_paragraph(f"Channel: {video_info['channel_title']}")
    document.add_paragraph(f"Published on: {video_info['upload_date_str']}")
    document.add_heading('Key Statistics', level=2)
    document.add_paragraph(f"Views: {video_info['views']:,}", style='List Bullet')
    document.add_paragraph(f"Likes: {video_info['likes']:,}", style='List Bullet')
    document.add_paragraph(f"Comments: {video_info['comments']:,}", style='List Bullet')
    document.add_heading('Tags', level=2)
    document.add_paragraph(', '.join(video_info['tags']) if video_info['tags'] else "No tags found.")
    document.add_heading('Description', level=2)
    document.add_paragraph(video_info['description'])
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    filename = f"{sanitize_filename(video_info.get('title', 'video'))}_analysis.docx"

    # --- यहाँ बदलाव किया गया है ---
    ascii_filename = filename.encode('ascii', 'ignore').decode('ascii', 'ignore')
    headers = {
        'Content-Disposition': 'attachment; filename*=UTF-8\'\'{}; filename="{}"'.format(quote(filename), ascii_filename)
    }

    return Response(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers=headers
    )

@analysis_bp.route('/analysis/export_video/txt/<string:video_id>')
@login_required
def export_single_video_to_txt(video_id):
    video_info = get_video_info_dict(video_id)
    if 'error' in video_info:
        flash(f"Could not export data: {video_info['error']}", "error")
        return redirect(url_for('analysis.video_analysis', video_id=video_id))

    text_content = f"--- Video Analysis for: {video_info['title']} ---\n\n"
    text_content += f"Channel: {video_info['channel_title']}\n"
    text_content += f"URL: https://www.youtube.com/watch?v={video_id}\n\n"
    text_content += "--- STATISTICS ---\n"
    text_content += f"Views: {video_info['views']:,}\n"
    text_content += f"Likes: {video_info['likes']:,}\n"
    text_content += f"Comments: {video_info['comments']:,}\n\n"
    text_content += "--- TAGS ---\n"
    text_content += f"{', '.join(video_info['tags']) if video_info['tags'] else 'No tags found.'}\n\n"
    text_content += "--- DESCRIPTION ---\n"
    text_content += video_info['description']

    buffer = BytesIO(text_content.encode('utf-8'))
    buffer.seek(0)
    
    filename = f"{sanitize_filename(video_info.get('title', 'video'))}_analysis.txt"

    # --- यहाँ बदलाव किया गया है ---
    ascii_filename = filename.encode('ascii', 'ignore').decode('ascii', 'ignore')
    headers = {
        'Content-Disposition': 'attachment; filename*=UTF-8\'\'{}; filename="{}"'.format(quote(filename), ascii_filename)
    }

    return Response(
        buffer,
        mimetype='text/plain',
        headers=headers
    )

@analysis_bp.route('/analysis/export/word/<string:channel_id>')
@login_required
def export_analysis_to_word(channel_id):
    channel_data = analyze_channel(channel_id)
    if 'error' in channel_data:
        flash(f"Could not generate report: {channel_data['error']}", "error")
        return redirect(url_for('analysis.deep_analysis', channel_id=channel_id))
        
    top_tags = get_most_used_tags(channel_id, video_limit=50)
    most_viewed_data = get_most_viewed_videos(channel_id, max_results=5) or {}
    recent_videos_data = get_latest_videos(channel_id, max_results=5) or {}
    most_viewed_videos = most_viewed_data.get('videos', [])
    recent_videos = recent_videos_data.get('videos', [])

    document = Document()
    document.add_heading('YouTube Channel Analysis Report', level=0)
    document.add_heading(f"Analysis for: {channel_data.get('Title', 'N/A')}", level=1)
    
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Subscribers'
    hdr_cells[1].text = 'Total Views'
    hdr_cells[2].text = 'Total Videos'
    hdr_cells[3].text = 'Channel ID'
    
    row_cells = table.add_row().cells
    row_cells[0].text = f"{int(channel_data.get('Subscribers', 0)):,}"
    row_cells[1].text = f"{int(channel_data.get('Total Views', 0)):,}"
    row_cells[2].text = str(channel_data.get('Video Count', 0))
    row_cells[3].text = channel_data.get('id', 'N/A')
    
    document.add_heading('Top Used Tags', level=2)
    if top_tags:
        for tag, count in top_tags:
            document.add_paragraph(f"{tag} (Used in {count} videos)", style='List Bullet')
    else:
        document.add_paragraph("No significant tags found.")

    document.add_heading('Top 5 Most Viewed Videos', level=2)
    if most_viewed_videos:
        for video in most_viewed_videos:
            p = document.add_paragraph()
            p.add_run(f"{video.get('title', 'N/A')}").bold = True
            p.add_run(f"\nViews: {video.get('view_count', 0):,}").italic = True
    else:
        document.add_paragraph("Could not fetch most viewed videos.")

    document.add_heading('Top 5 Recent Videos', level=2)
    if recent_videos:
        for video in recent_videos:
            p = document.add_paragraph()
            p.add_run(f"{video.get('title', 'N/A')}").bold = True
            if video.get('upload_date'):
                upload_date = datetime.fromisoformat(video['upload_date'].replace('Z', '')).strftime('%d %b, %Y')
                p.add_run(f"\nPublished on: {upload_date}").italic = True
    else:
        document.add_paragraph("Could not fetch recent videos.")
        
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"{sanitize_filename(channel_data.get('Title', 'channel'))}_report.docx"
    
    # --- यहाँ बदलाव किया गया है ---
    ascii_filename = filename.encode('ascii', 'ignore').decode('ascii', 'ignore')
    headers = {
        'Content-Disposition': 'attachment; filename*=UTF-8\'\'{}; filename="{}"'.format(quote(filename), ascii_filename)
    }
    
    return Response(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers=headers
    )